"""DOCX knowledge handler."""
from __future__ import annotations

import base64
import io
import logging
from typing import Any, Dict, Optional

from skills.knowledge import register_skill

logger = logging.getLogger(__name__)

MIME_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
]


async def extract(envelope: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    extra = envelope.get("extra", {})
    mimetype = (
        extra.get("doc_mimetype")
        or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ).lower()
    source_name = extra.get("doc_file_name") or "document.docx"
    raw_bytes = await _download_bytes(envelope)
    if raw_bytes is None:
        return None
    text = ""
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw_bytes))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            parts.append("[TABELA]")
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        text = "\n".join(parts)
    except Exception as exc:
        logger.warning("docx_handler extract failed: %s", exc)
        return None
    return {
        "text": text,
        "source_name": source_name,
        "mimetype": mimetype,
        "raw_size": len(raw_bytes),
    }


async def _download_bytes(envelope: Dict[str, Any]) -> Optional[bytes]:
    extra = envelope.get("extra", {})
    doc_b64 = extra.get("doc_base64", "")
    if doc_b64:
        try:
            return base64.b64decode(doc_b64)
        except Exception:
            pass
    message_id = envelope.get("message_id", "")
    instance = envelope.get("instance", "")
    remote_jid = extra.get("remote_jid", "")
    if message_id and instance:
        try:
            from core.evolution_client import get_base64_from_media_message

            result = await get_base64_from_media_message(
                instance=instance,
                message_id=message_id,
                remote_jid=remote_jid,
            )
            media_b64 = result.get("base64", "")
            if media_b64:
                return base64.b64decode(media_b64)
        except Exception as exc:
            logger.warning("docx_handler get_base64 failed: %s", exc)
    return None


async def persist(
    envelope: Dict[str, Any],
    extracted: Dict[str, Any],
    scope: str,
    metadata=None,
) -> Dict[str, Any]:
    phone = envelope.get("phone", "")
    metadata = metadata or {}
    text = extracted.get("text", "")
    source_name = extracted.get("source_name", "document.docx")
    mimetype = extracted.get("mimetype", MIME_TYPES[0])
    metadata = metadata or {}
    extra_metadata = {"mimetype": mimetype, "scope": scope, **metadata}
    if not text:
        return {"error": "no_text_extracted", "mimetype": mimetype}

    if scope == "group":
        try:
            from tools.group import index_group_document

            group_jid = envelope.get("extra", {}).get("remote_jid", "")
            if "@g.us" not in group_jid:
                return {"error": "group_jid_required"}
            result = await index_group_document(
                phone=phone,
                group_jid=group_jid,
                text=text,
                visibility="group",
                source_name=source_name,
            )
            return {
                "status": "rag_group",
                "index_result": result,
                "source_name": source_name,
                "scope": scope,
            }
        except Exception as exc:
            logger.warning("docx_handler group index failed: %s", exc)
            return {"error": "rag_index_failed", "detail": str(exc)}

    try:
        from core.rag import index_private_document

        result = await index_private_document(
            phone=phone,
            text_content=text,
            source_title=source_name,
            category="whatsapp_attachment",
            metadata=extra_metadata,
        )
        if result.get("error"):
            return {"error": "rag_index_failed", "detail": result.get("error")}
        return {
            "status": "rag_individual",
            "index_result": result,
            "source_name": source_name,
            "scope": scope,
        }
    except Exception as exc:
        logger.warning("docx_handler private index failed: %s", exc)
        return {"error": "rag_index_failed", "detail": str(exc)}


register_skill("docx", type("DOCXSkill", (), {
    "MIME_TYPES": MIME_TYPES,
    "extract": staticmethod(extract),
    "persist": staticmethod(persist),
})())

__all__ = ["extract", "persist", "MIME_TYPES"]
