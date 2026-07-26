"""Tests for google_drive tool (per-user OAuth, Fase D)."""
import io

import pytest
from unittest.mock import patch, MagicMock


@pytest.fixture
def mock_firestore(monkeypatch):
    fake_db = MagicMock()
    docs = [{"owner_phone": PHONE, "owner_uid": PHONE, "instance": "jennifer"}]

    class _FC:
        def where(self, *_a, **_k):
            return self

        def limit(self, *_a, **_k):
            return self

        def stream(self):
            items = list(docs)
            for item in items:
                captured = item
                yield MagicMock(to_dict=lambda c=captured: c, id=captured["instance"])

    fake_db.collection.return_value = _FC()
    monkeypatch.setattr("agent_loader._get_firestore_client", lambda: fake_db)
    monkeypatch.setattr("core.owner._get_firestore_client", lambda: fake_db)
    return fake_db


@pytest.fixture
def mock_drive_service(mock_firestore):
    with patch("tools.google_drive._get_service") as mock:
        service = MagicMock()
        mock.return_value = service
        yield service


PHONE = "5511966830020"


class TestSearchFiles:
    @pytest.mark.asyncio
    async def test_search_files(self, mock_drive_service):
        from tools.google_drive import search_files
        mock_drive_service.files().list().execute.return_value = {
            "files": [
                {"id": "f1", "name": "Ata 2026-07-13.md", "mimeType": "text/markdown"},
                {"id": "f2", "name": "Ata 2026-07-12.md", "mimeType": "text/markdown"},
            ]
        }
        result = await search_files(PHONE, "Ata", instance="jennifer")
        assert result["count"] == 2
        assert "Ata" in result["files"][0]["name"]


class TestUploadFile:
    @pytest.mark.asyncio
    async def test_upload_file(self, mock_drive_service):
        from tools.google_drive import upload_file
        mock_drive_service.files().create().execute.return_value = {
            "id": "new1", "name": "test.md", "mimeType": "text/markdown",
            "webViewLink": "https://drive.google.com/file/d/new1",
        }
        result = await upload_file(PHONE, "folder1", "test.md", "# Content", instance="jennifer")
        assert "file" in result
        assert result["file"]["name"] == "test.md"


class TestListFolder:
    @pytest.mark.asyncio
    async def test_list_folder(self, mock_drive_service):
        from tools.google_drive import list_folder
        mock_drive_service.files().list().execute.return_value = {
            "files": [
                {"id": "f1", "name": "doc.md", "mimeType": "text/markdown"},
                {"id": "f2", "name": "Subfolder", "mimeType": "application/vnd.google-apps.folder"},
            ]
        }
        result = await list_folder(PHONE, "parent_folder", instance="jennifer")
        assert result["count"] == 2


class TestCreateFolder:
    @pytest.mark.asyncio
    async def test_create_folder(self, mock_drive_service):
        from tools.google_drive import create_folder
        mock_drive_service.files().create().execute.return_value = {
            "id": "new_folder", "name": "NewFolder", "mimeType": "application/vnd.google-apps.folder",
        }
        result = await create_folder(PHONE, "NewFolder", parent_id="parent", instance="jennifer")
        assert "folder" in result
        assert result["folder"]["name"] == "NewFolder"


class TestFindOmnichannelAtas:
    @pytest.mark.asyncio
    async def test_find_existing(self, mock_drive_service):
        from tools.google_drive import find_omnichannel_atas_folder
        mock_drive_service.files().list().execute.side_effect = [
            {"files": [{"id": "omni_id"}]},
            {"files": [{"id": "atas_id"}]},
        ]
        result = await find_omnichannel_atas_folder(PHONE, instance="jennifer")
        assert result["folder_id"] == "atas_id"

    @pytest.mark.asyncio
    async def test_omnichannel_not_found(self, mock_drive_service):
        from tools.google_drive import find_omnichannel_atas_folder
        mock_drive_service.files().list().execute.return_value = {"files": []}
        result = await find_omnichannel_atas_folder(PHONE, instance="jennifer")
        assert "error" in result


class TestReadFileContent:
    """Tests for google_drive.read_file_content (PDF, DOCX, XLSX, Google Docs)."""

    @staticmethod
    def _make_downloader_side_effect(raw_bytes):
        """Build a MediaIoBaseDownload fake that writes raw_bytes into the real buf.
        Usage: patch('tools.google_drive.MediaIoBaseDownload', side_effect=side_effect)
        """
        def _side_effect(buf, request):
            dl = MagicMock()
            done_state = {"done": False}

            def _next_chunk():
                if not done_state["done"]:
                    buf.write(raw_bytes)
                    done_state["done"] = True
                return (None, done_state["done"])

            dl.next_chunk.side_effect = _next_chunk
            return dl
        return _side_effect

    @pytest.mark.asyncio
    async def test_read_pdf(self, mock_drive_service):
        from tools.google_drive import read_file_content
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(72, 72)
        buf = io.BytesIO()
        writer.write(buf)
        raw_pdf = buf.getvalue()

        mock_drive_service.files().get().execute.return_value = {
            "name": "Ata.pdf",
            "mimeType": "application/pdf",
            "size": str(len(raw_pdf)),
        }
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(raw_pdf)):
            result = await read_file_content(PHONE, "file_pdf_1", instance="jennifer")
        assert "file_id" in result
        assert result["parser"] == "pdf"
        assert result["file_name"] == "Ata.pdf"
        assert isinstance(result.get("content"), str)

    @pytest.mark.asyncio
    async def test_read_docx(self, mock_drive_service):
        from tools.google_drive import read_file_content
        from docx import Document

        doc = Document()
        doc.add_paragraph("Linha um do documento.")
        doc.add_paragraph("Segunda linha importante.")
        buf = io.BytesIO()
        doc.save(buf)
        raw_docx = buf.getvalue()

        mock_drive_service.files().get().execute.return_value = {
            "name": "Resumo.docx",
            "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "size": str(len(raw_docx)),
        }
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(raw_docx)):
            result = await read_file_content(PHONE, "file_docx_1", instance="jennifer")
        assert result["parser"] == "docx"
        assert "Linha um" in result["content"]
        assert "Segunda linha" in result["content"]

    @pytest.mark.asyncio
    async def test_read_xlsx_returns_ascii_table(self, mock_drive_service):
        from tools.google_drive import read_file_content
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Atas"
        ws.append(["Data", "Titulo", "Participantes"])
        ws.append(["2026-07-20", "Kickoff", "5"])
        ws.append(["2026-07-22", "Review", "3"])
        buf = io.BytesIO()
        wb.save(buf)
        raw_xlsx = buf.getvalue()

        mock_drive_service.files().get().execute.return_value = {
            "name": "Atas.xlsx",
            "mimeType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "size": str(len(raw_xlsx)),
        }
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(raw_xlsx)):
            result = await read_file_content(PHONE, "file_xlsx_1", instance="jennifer")
        assert result["parser"] == "xlsx"
        content = result["content"]
        assert ("|" in content) or ("Atas" in content)
        assert "Kickoff" in content

    @pytest.mark.asyncio
    async def test_read_google_doc_uses_export(self, mock_drive_service):
        from tools.google_drive import read_file_content

        mock_drive_service.files().get().execute.return_value = {
            "name": "Doc sem titulo",
            "mimeType": "application/vnd.google-apps.document",
            "size": "0",
        }
        raw_text = b"Conteudo exportado do Google Doc.\nLinha dois."
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(raw_text)) as mock_dl:
            result = await read_file_content(PHONE, "file_gdoc_1", instance="jennifer")
        assert result["parser"] == "google_doc"
        assert "Conteudo exportado" in result["content"]
        mock_dl.assert_called()

    @pytest.mark.asyncio
    async def test_read_text_file(self, mock_drive_service):
        from tools.google_drive import read_file_content

        raw = b"linha 1\nlinha 2\nlinha 3\n"
        mock_drive_service.files().get().execute.return_value = {
            "name": "notas.md",
            "mimeType": "text/markdown",
            "size": str(len(raw)),
        }
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(raw)):
            result = await read_file_content(PHONE, "file_txt_1", instance="jennifer")
        assert result["parser"] == "text"
        assert "linha 2" in result["content"]

    @pytest.mark.asyncio
    async def test_unsupported_mime_returns_error(self, mock_drive_service):
        from tools.google_drive import read_file_content

        mock_drive_service.files().get().execute.return_value = {
            "name": "video.mp4",
            "mimeType": "video/mp4",
            "size": "1000",
        }
        result = await read_file_content(PHONE, "file_mp4_1", instance="jennifer")
        assert result.get("error") == "unsupported_mime_type"
        assert result["mime_type"] == "video/mp4"

    @pytest.mark.asyncio
    async def test_empty_file_id_returns_error(self, mock_drive_service):
        from tools.google_drive import read_file_content

        result = await read_file_content(PHONE, "", instance="jennifer")
        assert result.get("error") == "file_id_required"

    @pytest.mark.asyncio
    async def test_content_truncated_above_limit(self, mock_drive_service):
        """Quando o conteudo extraido excede 12k chars, deve truncar e marcar truncated=True."""
        from tools.google_drive import read_file_content

        big_text = ("X" * 15000).encode("utf-8")
        mock_drive_service.files().get().execute.return_value = {
            "name": "grande.txt",
            "mimeType": "text/plain",
            "size": str(len(big_text)),
        }
        with patch("tools.google_drive.MediaIoBaseDownload",
                    side_effect=self._make_downloader_side_effect(big_text)):
            result = await read_file_content(PHONE, "file_big_1", instance="jennifer")
        assert result["truncated"] is True
        assert "truncated" in result["content"].lower()


class TestPerUserOAuth:
    def test_get_service_uses_per_user_oauth(self):
        from tools import google_drive
        from core import oauth_per_user

        google_drive._drive_services.clear()
        with patch.object(google_drive, "build", return_value="service-mock") as mock_build:
            with patch.object(oauth_per_user, "get_user_credentials", return_value=MagicMock()) as mock_user_creds:
                service = google_drive._get_service(PHONE)
                assert mock_user_creds.called
                assert mock_user_creds.call_args.args[0] == PHONE
                assert service == "service-mock"
                assert mock_build.called
                assert mock_build.call_args.args[0] == "drive"

    def test_get_credentials_requires_phone(self):
        from tools import google_drive

        with pytest.raises(RuntimeError, match="phone_required_for_drive_oauth"):
            google_drive._get_credentials("")

    def test_get_credentials_requires_user_setup(self):
        from tools import google_drive
        from core import oauth_per_user

        with patch.object(oauth_per_user, "get_user_credentials", return_value=None):
            with pytest.raises(RuntimeError, match="user_google_oauth_required"):
                google_drive._get_credentials(PHONE)
