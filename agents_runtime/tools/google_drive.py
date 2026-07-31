"""Google Drive tools - 4 functions.

Auth: per-user OAuth via core.oauth_per_user.get_user_credentials.
The phone parameter is mandatory (Fase D); the global GOOGLE_OAUTH_TOKEN
fallback was removed.

Owner-only: Drive access is restricted to the phone bound to the Evolution
instance. Any other phone is denied with an ``owner_only_capability`` error.
"""
import functools
import logging
from typing import Optional, Dict, Any

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
import io

logger = logging.getLogger(__name__)

SCOPES = [
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/drive.readonly",
]
_OMNICHANNEL_FOLDER_NAME = "Omnichannel"
_ATAS_SUBFOLDER = "Atas"

_drive_services: Dict[str, Any] = {}


def _get_credentials(phone: str) -> Credentials:
    """Load Google OAuth credentials for the given user (per-user, Fase D)."""
    if not phone:
        raise RuntimeError("phone_required_for_drive_oauth")
    from core.oauth_per_user import get_user_credentials

    creds = get_user_credentials(phone)
    if creds is None:
        raise RuntimeError("user_google_oauth_required")
    return creds


def _get_service(phone: str):
    """Get or build Drive API service (cached per user)."""
    cache_key = phone
    if cache_key not in _drive_services:
        creds = _get_credentials(phone)
        _drive_services[cache_key] = build("drive", "v3", credentials=creds, cache_discovery=False)
    return _drive_services[cache_key]


def _format_file(file: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": file.get("id"),
        "name": file.get("name"),
        "mime_type": file.get("mimeType"),
        "modified": file.get("modifiedTime"),
        "size": file.get("size"),
        "web_view_link": file.get("webViewLink"),
    }


def _owner_guard(capability: str):
    from core.owner import deny_if_not_owner, resolve_owner
    from core.owner_guard import check_folder_permission, post_filter_tool_result

    def decorator(func):
        @functools.wraps(func)
        async def wrapper(*args, **kwargs):
            phone = kwargs.get("phone")
            if not phone and args:
                phone = args[0]
            phone = str(phone or "")
            instance = str(kwargs.get("instance", "") or kwargs.get("_instance", ""))
            resolution = resolve_owner(instance, fallback_phone=phone)
            denial = deny_if_not_owner(resolution, phone, capability)
            if denial is not None:
                return denial
            fp_denial = check_folder_permission(phone, capability, kwargs)
            if fp_denial is not None:
                return fp_denial
            result = await func(*args, **kwargs)
            return await post_filter_tool_result(phone, capability, result, kwargs)
        return wrapper
    return decorator


@_owner_guard("drive.search")
async def search_files(
    phone: str,
    query: str,
    folder_id: Optional[str] = None,
    mime_type: Optional[str] = None,
    max_results: int = 20,
    instance: str = "",
) -> Dict[str, Any]:
    """Search for files in Drive.

    Args:
        phone: User phone for per-user OAuth token (mandatory, Fase D).
        query: Search query (Google Drive query syntax)
        folder_id: Restrict search to folder
        mime_type: Filter by MIME type
        max_results: Max results

    Returns:
        {"files": [...], "count": int}
    """
    try:
        service = _get_service(phone)
        q_parts = [f"name contains '{query}'" if query else ""]
        if folder_id:
            q_parts.append(f"'{folder_id}' in parents")
        if mime_type:
            q_parts.append(f"mimeType='{mime_type}'")
        q = " and ".join(p for p in q_parts if p)

        result = service.files().list(
            q=q if q else None,
            pageSize=max_results,
            fields="files(id, name, mimeType, modifiedTime, size, webViewLink)",
        ).execute()
        files = [_format_file(f) for f in result.get("files", [])]
        return {"files": files, "count": len(files)}
    except HttpError as e:
        logger.error(f"Drive search_files error: {e}")
        return {"files": [], "count": 0, "error": str(e)}


@_owner_guard("drive.upload")
async def upload_file(
    phone: str,
    folder_id: str,
    filename: str,
    content: str,
    mime_type: str = "text/plain",
    instance: str = "",
) -> Dict[str, Any]:
    """Upload a file to Drive folder.

    Args:
        phone: User phone for per-user OAuth token (mandatory, Fase D).
        folder_id: Destination folder ID
        filename: Name of the file
        content: File content (string)
        mime_type: MIME type (default text/plain)

    Returns:
        {"file": {...}} or {"error": str}
    """
    try:
        service = _get_service(phone)
        file_metadata = {
            "name": filename,
            "parents": [folder_id],
        }
        media = MediaIoBaseUpload(
            io.BytesIO(content.encode("utf-8")),
            mimetype=mime_type,
            resumable=True,
        )
        uploaded = service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id, name, mimeType, modifiedTime, size, webViewLink",
        ).execute()
        return {"file": _format_file(uploaded)}
    except HttpError as e:
        logger.error(f"Drive upload_file error: {e}")
        return {"error": str(e)}


@_owner_guard("drive.list")
async def list_folder(
    phone: str,
    folder_id: str = "root",
    max_results: int = 50,
    instance: str = "",
) -> Dict[str, Any]:
    """List contents of a folder.

    Args:
        phone: User phone for per-user OAuth token (mandatory, Fase D).
        folder_id: Folder ID (default: root)
        max_results: Max results

    Returns:
        {"files": [...], "count": int}
    """
    try:
        service = _get_service(phone)
        if folder_id == "root":
            query = "'root' in parents"
        else:
            query = f"'{folder_id}' in parents"

        result = service.files().list(
            q=query,
            pageSize=max_results,
            fields="files(id, name, mimeType, modifiedTime, size, webViewLink)",
        ).execute()
        files = [_format_file(f) for f in result.get("files", [])]
        return {"files": files, "count": len(files)}
    except HttpError as e:
        logger.error(f"Drive list_folder error: {e}")
        return {"files": [], "count": 0, "error": str(e)}


@_owner_guard("drive.create_folder")
async def create_folder(
    phone: str,
    name: str,
    parent_id: Optional[str] = None,
    instance: str = "",
) -> Dict[str, Any]:
    """Create a folder.

    Args:
        phone: User phone for per-user OAuth token (mandatory, Fase D).
        name: Folder name
        parent_id: Parent folder ID (None for root)

    Returns:
        {"folder": {...}} or {"error": str}
    """
    try:
        service = _get_service(phone)
        metadata = {
            "name": name,
            "mimeType": "application/vnd.google-apps.folder",
        }
        if parent_id:
            metadata["parents"] = [parent_id]
        folder = service.files().create(body=metadata, fields="id, name, mimeType").execute()
        return {"folder": _format_file(folder)}
    except HttpError as e:
        logger.error(f"Drive create_folder error: {e}")
        return {"error": str(e)}


@_owner_guard("drive.find_omnichannel_atas")
async def find_omnichannel_atas_folder(phone: str, instance: str = "") -> Dict[str, Any]:
    """Find the Omnichannel/Atas/ folder, creating if missing.

    Args:
        phone: User phone for per-user OAuth token (mandatory, Fase D).

    Returns:
        {"folder_id": "..."} or {"error": str}
    """
    try:
        service = _get_service(phone)
        results = service.files().list(
            q=f"name='{_OMNICHANNEL_FOLDER_NAME}' and mimeType='application/vnd.google-apps.folder'",
            fields="files(id, name)",
        ).execute()
        files = results.get("files", [])
        if not files:
            return {"error": f"{_OMNICHANNEL_FOLDER_NAME} folder not found"}
        omnichannel_id = files[0]["id"]

        atas_results = service.files().list(
            q=f"name='{_ATAS_SUBFOLDER}' and mimeType='application/vnd.google-apps.folder' and '{omnichannel_id}' in parents",
            fields="files(id, name)",
        ).execute()
        atas_files = atas_results.get("files", [])
        if atas_files:
            return {"folder_id": atas_files[0]["id"]}

        created = await create_folder(phone, _ATAS_SUBFOLDER, parent_id=omnichannel_id)
        if "folder" in created:
            return {"folder_id": created["folder"]["id"]}
        return {"error": "failed to create Atas folder"}
    except HttpError as e:
        logger.error(f"Drive find_omnichannel_atas_folder error: {e}")


_MAX_EXTRACT_CHARS = 12000  # ~2000 palavras / ~8 paginas, cobrem 95% dos casos

_GOOGLE_DOCS_MIME_PREFIX = "application/vnd.google-apps"
_TEXT_PLAIN = "text/plain"
_TEXT_CSV = "text/csv"


def _parse_pdf(raw: bytes) -> str:
    from core.pdf_extract import parse_pdf_robust

    return parse_pdf_robust(raw)


def _parse_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_xlsx(raw: bytes, max_rows: int = 50, max_cols: int = 6) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    chunks = []
    for sheet_name in workbook.sheetnames[:3]:
        sheet = workbook[sheet_name]
        chunks.append(f"--- Sheet: {sheet_name} ---")
        for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if idx > max_rows:
                chunks.append(f"... ({max_rows}+ rows truncated)")
                break
            values = [
                "" if v is None else str(v)[:80]
                for v in row[:max_cols]
            ]
            if any(v for v in values):
                chunks.append(" | ".join(values))
    workbook.close()
    return "\n".join(chunks)


def _format_xlsx_as_table(raw: bytes, max_rows: int = 30, max_cols: int = 5) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    sheets = []
    for sheet_name in workbook.sheetnames[:2]:
        sheet = workbook[sheet_name]
        rows = []
        for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if idx > max_rows:
                rows.append(f"... ({max_rows}+ rows truncated)")
                break
            values = [
                "" if v is None else str(v)[:40].replace("\n", " ")
                for v in row[:max_cols]
            ]
            rows.append(values)
        if rows:
            sheets.append((sheet_name, rows))
    workbook.close()
    if not sheets:
        return ""
    out = []
    for sheet_name, rows in sheets:
        out.append(f"*{sheet_name}*")
        if not rows:
            out.append("(empty)")
            continue
        width = max(len(r) for r in rows)
        widths = [max(len(str(r[i])) if i < len(r) else 0 for r in rows) for i in range(width)]
        widths = [min(w, 32) for w in widths]
        for i, r in enumerate(rows):
            cells = [str(r[j]) if j < len(r) else "" for j in range(width)]
            line = " | ".join(c.ljust(widths[j]) for j, c in enumerate(cells))
            out.append(line)
            if i == 0:
                out.append("-+-".join("-" * w for w in widths))
        out.append("")
    return "\n".join(out)


def _truncate(text: str, limit: int = _MAX_EXTRACT_CHARS) -> tuple:
    if len(text) <= limit:
        return text, False
    return text[:limit] + f"\n... [truncated at {limit} chars]", True


@_owner_guard("drive.read_file")
async def read_file_content(
    phone: str,
    file_id: str,
    instance: str = "",
) -> Dict[str, Any]:
    """Download and extract text content from a Google Drive file.

    Supports:
    - PDF (application/pdf) — parsed via pypdf
    - DOCX (Office Open XML Word) — parsed via python-docx
    - XLSX (Office Open XML Spreadsheet) — formatted as WhatsApp-friendly
      ASCII table when possible; falls back to pipe-delimited rows
    - Plain text / CSV (text/*) — decoded as UTF-8
    - Google Docs / Sheets / Slides (application/vnd.google-apps.*) —
      exported to text/plain or text/csv via the export endpoint

    Returns:
        {
            "file_id": str,
            "file_name": str,
            "mime_type": str,
            "size": int,
            "content": str,        # extracted text, truncated to 12k chars
            "truncated": bool,     # True if extraction hit the cap
            "parser": str          # "pdf"|"docx"|"xlsx"|"text"|"google_doc"|...
        }
    """
    if not file_id:
        return {"error": "file_id_required"}
    try:
        service = _get_service(phone)
        meta = service.files().get(
            fileId=file_id, fields="name,mimeType,size"
        ).execute()
        mime_type = meta.get("mimeType", "")
        file_name = meta.get("name", "")
        size = int(meta.get("size") or 0)

        if mime_type == "application/pdf":
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            text = _parse_pdf(raw)
            parser = "pdf"
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            text = _parse_docx(raw)
            parser = "docx"
        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            formatted = _format_xlsx_as_table(raw)
            text = formatted if formatted else _parse_xlsx(raw)
            parser = "xlsx"
        elif mime_type.startswith(f"{_GOOGLE_DOCS_MIME_PREFIX}."):
            export_mime = _TEXT_CSV if "spreadsheet" in mime_type else _TEXT_PLAIN
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(
                buf,
                service.files().export(fileId=file_id, mimeType=export_mime),
            )
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            text = raw.decode("utf-8", errors="replace")
            parser = "google_doc" if "document" in mime_type else (
                "google_sheet" if "spreadsheet" in mime_type else "google_slides"
            )
        elif mime_type.startswith("text/"):
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            text = raw.decode("utf-8", errors="replace")
            parser = "text"
        else:
            return {
                "error": "unsupported_mime_type",
                "file_id": file_id,
                "file_name": file_name,
                "mime_type": mime_type,
            }

        text, truncated = _truncate(text)
        return {
            "file_id": file_id,
            "file_name": file_name,
            "mime_type": mime_type,
            "size": size,
            "content": text,
            "truncated": truncated,
            "parser": parser,
        }
    except HttpError as e:
        logger.error(f"Drive read_file_content error: {e}")
        return {"error": f"drive_api:{e}", "file_id": file_id}
    except Exception as e:
        logger.error(f"Drive read_file_content unexpected error: {e}")
        return {"error": f"parse:{type(e).__name__}:{e}", "file_id": file_id}


_FOLDER_NAME_SYNONYMS = {
    "ata": ["ata", "atas", "minuta", "minutas", "reuniao", "reunioes", "meeting", "minutes"],
    "relatorio": ["relatorio", "relatorios", "report", "reports"],
    "projeto": ["projeto", "projetos", "project", "projects"],
    "orcamento": ["orcamento", "orcamentos", "budget", "custo", "custos", "financeiro"],
}

_MAX_DEEP_SEARCH_RESULTS = 50
_DEFAULT_MAX_DEPTH = 3


def _match_folder_name(folder_name: str, query: str) -> bool:
    fn = folder_name.lower().strip()
    q = query.lower().strip()
    if q in fn:
        return True
    for base, synonyms in _FOLDER_NAME_SYNONYMS.items():
        if base in q or any(s in q for s in synonyms):
            if any(s in fn for s in synonyms):
                return True
    return False


@_owner_guard("drive.deep_search")
async def deep_search_drive(
    phone: str,
    query: str,
    parent_folder_id: str = "root",
    max_depth: int = _DEFAULT_MAX_DEPTH,
    max_results: int = _MAX_DEEP_SEARCH_RESULTS,
    include_shared_drives: bool = True,
    instance: str = "",
) -> Dict[str, Any]:
    """Recursive search across Drive folders and subfolders.

    Uses BFS to scan folder tree, matching both file names and folder
    names semantically. Supports shared drives.

    Args:
        phone: User phone for per-user OAuth token.
        query: Search query (matched against file AND folder names).
        parent_folder_id: Starting folder (default: "root" = all drives).
        max_depth: Maximum recursion depth (default 3).
        max_results: Maximum total results (default 50).
        include_shared_drives: Search shared drives too (default True).
        instance: Evolution instance name.

    Returns:
        {"files": [...], "count": int, "scanned_folders": int, "max_depth_reached": int}
    """
    if not query:
        return {"files": [], "count": 0, "error": "query_required"}

    try:
        service = _get_service(phone)
        queue = [(parent_folder_id, 0)]
        visited = set()
        all_files = []
        scanned = 0
        max_depth_reached = 0

        while queue and len(all_files) < max_results:
            current_id, depth = queue.pop(0)
            if current_id in visited:
                continue
            visited.add(current_id)
            max_depth_reached = max(max_depth_reached, depth)
            scanned += 1

            extra_params = {}
            if include_shared_drives and current_id == "root":
                extra_params["corpora"] = "allDrives"
                extra_params["includeItemsFromAllDrives"] = True
                extra_params["supportsAllDrives"] = True

            result = service.files().list(
                q=f"'{current_id}' in parents and trashed = false",
                pageSize=min(max_results, 100),
                fields="files(id, name, mimeType, modifiedTime, size, webViewLink)",
                orderBy="modifiedTime desc",
                **extra_params,
            ).execute()
            page_files = result.get("files", [])

            for f in page_files:
                fn = (f.get("name") or "").lower()
                qt = query.lower()
                if qt in fn or _match_folder_name(fn, query):
                    all_files.append(_format_file(f))

            if len(all_files) >= max_results:
                break

            if depth < max_depth:
                subfolders = [
                    f for f in page_files
                    if f.get("mimeType") == "application/vnd.google-apps.folder"
                ]
                for sf in subfolders:
                    queue.append((sf["id"], depth + 1))

        return {
            "files": all_files[:max_results],
            "count": len(all_files[:max_results]),
            "scanned_folders": scanned,
            "max_depth_reached": max_depth_reached,
            "query": query,
        }
    except HttpError as e:
        logger.error(f"Drive deep_search_drive error: {e}")
        return {"files": [], "count": 0, "error": str(e)}
